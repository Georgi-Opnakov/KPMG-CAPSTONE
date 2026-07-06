from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
README_PATH = BASE_DIR / "README.md"
OUTPUT_PATH = BASE_DIR / "README.docx"


def set_run_font(run, name: str) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_code_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "3")
        element.set(qn("w:color"), "D9E2EF")
        borders.append(element)


def add_inline_markdown(paragraph, text: str, *, code_font: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_run_font(run, "Calibri")
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 58, 95)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Consolas" if code_font else "Calibri")
            if code_font:
                run.font.size = Pt(9.5)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for idx, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.10)
        paragraph.paragraph_format.space_before = Pt(2 if idx else 6)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(paragraph, "F5F7FA")
        if idx == 0:
            set_code_border(paragraph)
        add_inline_markdown(paragraph, line or " ", code_font=True)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document() -> None:
    markdown = README_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Lilly Airbnb Intelligent Advisor")
    set_run_font(run, "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("README converted from Markdown")
    set_run_font(subtitle_run, "Calibri")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(91, 103, 112)

    in_code = False
    code_lines: list[str] = []
    skip_first_h1 = True

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            if skip_first_h1:
                skip_first_h1 = False
                continue
            paragraph = doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:].strip())
            continue

        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, stripped)

    if code_lines:
        add_code_block(doc, code_lines)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Lilly Airbnb Intelligent Advisor")
    set_run_font(footer_run, "Calibri")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(91, 103, 112)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
